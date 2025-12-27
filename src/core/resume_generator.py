# %%
# Import Required Libraries
from anthropic import Anthropic
import json
import yaml
import re
from datetime import datetime
import os
from dotenv import load_dotenv
import textwrap

# Data validation & handling
from typing import List, Dict, Any, Optional
from pydantic import BaseModel, Field, model_validator
from pathlib import Path

# Word Document
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# PDF
from docx2pdf import convert
import pythoncom

# BrightData script (for job scraping)
from src.integrations.brightdata import extract_job

# Vector store and embeddings
from src.core.embeddings import OpenAIEmbeddings
from src.core.vector_store import QdrantVectorStore
from collections import defaultdict

# %%
# Load Environment Variables
load_dotenv()

key = os.getenv("CLAUDE_API_KEY")
if not key:
    raise ValueError("CLAUDE_API_KEY environment variable is required.")

model_name = "claude-haiku-4-5"

# %%
# Job Analysis Functions
def create_analysis_prompt(job_title: str, company: str, job_description: str) -> str:
    prompt_analysis = f"""
    You are an expert recruiter in the data analytics, data science, and AI engineering fields. Output in JSON format only.
    Your task is to extract the following information from the job details provided:

    1. job_summary: Create a concise 250-word summary of the job posting that captures the essence of the role, key responsibilities, and what the company is looking for. This summary should be clear, professional, and suitable for use in a cover letter.
    2. technical_skills: An array of technical skills required or preferred for the role. Look for specific technologies, programming languages, tools, or technical methodologies mentioned.
    3. soft_skills: An array of soft skills or personal attributes required or preferred for the role. Identify personal attributes, interpersonal skills, or work style preferences described.
    4. keywords: An array of keywords relevant to the role, including industry-specific terms, tools, or methodologies. Extract key terms that are frequently mentioned or seem particularly important to the role or industry.

    Here is the job information you need to analyze:

    <job_description>
    {job_description}
    </job_description>

    <job_title>
    {job_title}
    </job_title>

    <company>
    {company}
    </company>

    Output the information into a JSON object only with the following structure:

    {{
        "job_summary": "A 250-word summary of the job posting",
        "technical_skills": ["skill1", "skill2", "skill3"],
        "soft_skills": ["skill1", "skill2", "skill3"],
        "keywords": ["keyword1", "keyword2", "keyword3"]
    }}

    Each array should contain at least 3 items if possible and a maximum of 10, but don't include irrelevant information just to meet this number.

    """
    return prompt_analysis

# Job Analysis Response Data Validation
class JobAnalysisResponse(BaseModel):
    """Model for the job analysis response structure."""
    technical_skills: List[str] = Field(default_factory=list)
    soft_skills: List[str] = Field(default_factory=list)
    keywords: List[str] = Field(default_factory=list)
    job_summary: str = Field(default="", description="250-word summary of the job posting")
    
    @model_validator(mode='before')
    @classmethod
    def extract_from_text(cls, data):
        """Extract JSON from text if the input is not already a dict."""
        if isinstance(data, dict):
            return data
            
        if isinstance(data, str):
            # Try direct JSON parsing first
            try:
                return json.loads(data)
            except json.JSONDecodeError:
                # Extract JSON using regex if direct parsing fails
                try:
                    json_pattern = r'(\{[\s\S]*\})'
                    match = re.search(json_pattern, data)
                    if match:
                        return json.loads(match.group(1))
                except:
                    pass
                
                # If we can't extract valid JSON, try to create a structured response
                # by looking for keyword patterns
                skills_pattern = {
                    "job_summary": r'"?job_summary"?\s*:\s*"([^"]*)"',
                    "technical_skills": r'"?technical_skills"?\s*:\s*\[(.*?)\]',
                    "soft_skills": r'"?soft_skills"?\s*:\s*\[(.*?)\]',
                    "keywords": r'"?keywords"?\s*:\s*\[(.*?)\]'
                }
                
                extracted_data = {}
                for key, pattern in skills_pattern.items():
                    match = re.search(pattern, data, re.DOTALL)
                    if match:
                        if key == "job_summary":
                            # Extract string value for job_summary
                            extracted_data[key] = match.group(1)
                        else:
                            # Extract the items in the array
                            items_text = match.group(1)
                            # Parse the items, handling quoted strings
                            items = []
                            for item in re.findall(r'"([^"]*)"', items_text):
                                items.append(item)
                            if items:
                                extracted_data[key] = items
                
                if extracted_data:
                    return extracted_data
                    
                # If all else fails, print a warning and return empty lists/string
                print(f"Failed to parse analysis response: {data[:100]}...")
                return {
                    "technical_skills": [],
                    "soft_skills": [],
                    "keywords": [],
                    "job_summary": ""
                }
        
        return data
    

def claude_analysis(api_key: str, prompt_analysis: str) -> Dict[str, Any]:
    client = Anthropic(api_key=api_key)
    response = client.messages.create(
        model=model_name,
        max_tokens=2048,
        temperature=0.5,
        messages=[{"role": "user", "content": prompt_analysis}]
    )
    # Use Pydantic model to validate and extract the response
    response_text = response.content[0].text
    try:
        analysis_data = JobAnalysisResponse.model_validate(response_text)
        analysis_dict = analysis_data.model_dump()
        
        # Print formatted job analysis result
        print(f"\n{'='*80}")
        print("📊 JOB ANALYSIS RESULT")
        print(f"{'='*80}")
        print(json.dumps(analysis_dict, indent=2))
        print(f"{'='*80}\n")
        
        return analysis_dict
    except Exception as e:
        print(f"Error parsing response with Pydantic: {e}")
        # Fallback to simpler parsing if Pydantic validation fails
        try:
            return json.loads(response_text)
        except:
            print(f"JSON parsing failed. Raw response: {response_text[:150]}...")
            # Last resort: return empty data structure
            return {
                "technical_skills": [],
                "soft_skills": [],
                "keywords": [],
                "job_summary": ""
            }

def format_claude_analysis_response(skills_dict):
    formatted_text = ""    
    for category, items in skills_dict.items():
        # Convert category from snake_case to Title Case
        category_title = ' '.join(word.capitalize() for word in category.split('_'))
        # Section header
        formatted_text += f"## {category_title}\n"
        # Add bullet points for each item
        for item in items:
            formatted_text += f"• {item}\n"
        # Add extra line between sections
        formatted_text += "\n"
    return formatted_text

# %%
# Resume Generation Functions

# def load_resume_yaml(file_path: str) -> Dict[Any, Any]:
#     with open(file_path, 'r') as file:
#         resume_data = yaml.safe_load(file)
#     return resume_data


def retrieve_resume_context(
    job_title: Optional[str] = None, 
    company: Optional[str] = None, 
    job_description: Optional[str] = None, 
    top_k_achievement_pool: int = 100,
    top_k_achievements_per_job: int = 3,
    top_k_jobs: int = 4
) -> Dict[str, Any]:
    """
    Retrieve relevant resume context using RAG from vector database.
    
    This function queries the Qdrant vector store to retrieve the most relevant
    resume information based on job requirements. It supports two modes:
    1. Full retrieval (when job info is None): Returns all resume data
    2. Filtered retrieval (when job info is provided): Returns only relevant chunks
    
    Strategy:
    - Retrieves a large pool of achievements (top_k_achievement_pool)
    - Groups achievements by job (company + position + dates)
    - For each job, keeps only top N most relevant achievements (top_k_achievements_per_job)
    - Ranks jobs by average similarity score of their top achievements
    - Selects top N most relevant jobs (top_k_jobs)
    - Returns selected jobs with their top achievements (max 3 per job)
    
    Parameters
    ----------
    job_title : str, optional
        Job title to search for relevant experience.
    company : str, optional
        Company name (context for search).
    job_description : str, optional
        Full job description text for semantic search.
    top_k_achievement_pool : int, optional
        Size of achievement pool to retrieve from vector store (default: 100).
    top_k_achievements_per_job : int, optional
        Maximum achievements to keep per job (default: 3).
    top_k_jobs : int, optional
        Number of most relevant jobs to include (default: 4).
    
    Returns
    -------
    Dict[str, Any]
        Reconstructed resume data matching YAML structure:
        {
            "personal_information": {...},
            "professional_summary": {...},
            "work_experience": [{company, position, dates, achievements}, ...],
            "education": [{degree, institution, dates}, ...],
            "continuing_studies": [{name, institution, completion_date}, ...],
            "skills": {programming_languages: [...], ...}
        }
    
    Notes
    -----
    - Each job limited to max top_k_achievements_per_job achievements
    - Jobs with fewer achievements include all available
    - Ranks jobs by average similarity score of their top achievements
    - Uses semantic search on full job information for better matching
    
    Examples
    --------
    >>> # Full retrieval (initial pass)
    >>> resume = retrieve_resume_context()
    >>> 
    >>> # Filtered retrieval using job information
    >>> resume = retrieve_resume_context(
    ...     job_title="Data Scientist",
    ...     company="Acme Corp",
    ...     job_description="We need Python expertise for ETL pipelines...",
    ...     top_k_achievement_pool=100,
    ...     top_k_achievements_per_job=3,
    ...     top_k_jobs=4
    ... )
    """
    # Debug: Print parameters received
    print(f"\n🔍 DEBUG - retrieve_resume_context called with:")
    print(f"   job_title: {job_title}")
    print(f"   company: {company}")
    print(f"   job_description: {job_description[:100] if job_description else None}...")
    print()
    
    # Initialize embeddings and vector store
    embedder = OpenAIEmbeddings()
    store = QdrantVectorStore()
    
    # Construct query for filtered retrieval using full job information
    if job_title or company or job_description:
        query_parts = []
        if job_title:
            query_parts.append(job_title)
        if company:
            query_parts.append(company)
        if job_description:
            query_parts.append(job_description)
        
        query_text = ' '.join(query_parts)
        query_vector = embedder.embed_query(query_text)
        use_filtering = True
    else:
        # For full retrieval, use a generic query
        query_text = "data science analytics machine learning"
        query_vector = embedder.embed_query(query_text)
        use_filtering = False
    
    # Retrieve different sections with appropriate limits
    
    # Work experience achievements (filtered by relevance)
    work_results = store.search(
        query_vector=query_vector,
        top_k=top_k_achievement_pool if use_filtering else 100,  # Get large pool for filtering
        section_filter="work_experience"
    )
    
    # Education (get all)
    education_results = store.search(
        query_vector=query_vector,
        top_k=10,
        section_filter="education"
    )
    
    # Skills (get all)
    skills_results = store.search(
        query_vector=query_vector,
        top_k=20,
        section_filter="skills"
    )
    
    # Continuing studies
    continuing_results = store.search(
        query_vector=query_vector,
        top_k=10 if use_filtering else 100,
        section_filter="continuing_studies"
    )
    
    # Personal info (get the one entry)
    personal_results = store.search(
        query_vector=query_vector,
        top_k=1,
        section_filter="personal_info"
    )
    
    # Professional summary (get the one entry)
    summary_results = store.search(
        query_vector=query_vector,
        top_k=1,
        section_filter="professional_summary"
    )
    
    # Reconstruct resume structure
    resume_data = {
        "personal_information": {},
        "professional_summary": {"personality_traits": []},
        "work_experience": [],
        "education": {"degrees": [], "continuing_studies": []},
        "skills": {}
    }
    
    # Parse personal information
    if personal_results:
        # Extract from content (would need parsing, for now use metadata)
        content = personal_results[0]['content']
        lines = content.split('\n')
        for line in lines:
            if 'Email' in line:
                email_match = re.search(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', line)
                if email_match:
                    resume_data['personal_information']['email'] = email_match.group(1)
            if 'GitHub' in line:
                github_match = re.search(r'https?://github\.com/[^\s|]+', line)
                if github_match:
                    resume_data['personal_information']['github'] = github_match.group(0)
            if 'LinkedIn' in line:
                linkedin_match = re.search(r'https?://[^\s|]+linkedin[^\s|]+', line)
                if linkedin_match:
                    resume_data['personal_information']['linkedin'] = linkedin_match.group(0)
            if 'Languages' in line:
                lang_match = re.search(r'Languages\*\*: (.+)', line)
                if lang_match:
                    resume_data['personal_information']['languages'] = lang_match.group(1).strip()
            if 'Industry Experience' in line:
                ind_match = re.search(r'Industry Experience\*\*: (.+)', line)
                if ind_match:
                    resume_data['personal_information']['Industry Experience'] = ind_match.group(1).strip()
        
        # Extract name from first line
        name_match = re.match(r'# (.+)', lines[0])
        if name_match:
            resume_data['personal_information']['full_name'] = name_match.group(1).strip()
    
    # Professional summary
    if summary_results:
        content = summary_results[0]['content']
        # Split into traits (sentences)
        sentences = [s.strip() for s in content.split('.') if s.strip()]
        resume_data['professional_summary']['personality_traits'] = sentences
    
    # Group work experience achievements by job
    job_groups = defaultdict(lambda: {"achievements": [], "scores": [], "metadata": {}})
    for result in work_results:
        metadata = result['metadata']
        job_key = (
            metadata.get('company', ''),
            metadata.get('position', ''),
            metadata.get('start_date', ''),
            metadata.get('end_date', '')
        )
        
        # Use original achievement text from metadata (backward compatible fallback to content)
        achievement_text = result['metadata'].get('achievement_text', result['content'])
        job_groups[job_key]["achievements"].append(achievement_text)
        job_groups[job_key]["scores"].append(result['score'])
        job_groups[job_key]["metadata"] = {
            "location": metadata.get('location', ''),
            "industry": metadata.get('industry', '')
        }
    
    # For each job, keep only top N achievements by similarity score
    for job_key, job_data in job_groups.items():
        # Zip achievements with their scores for sorting
        achievement_score_pairs = list(zip(job_data["achievements"], job_data["scores"]))
        # Sort by score (descending)
        achievement_score_pairs.sort(key=lambda x: x[1], reverse=True)
        # Keep only top N (max top_k_achievements_per_job)
        top_pairs = achievement_score_pairs[:top_k_achievements_per_job]
        # Unzip back into separate lists
        job_data["achievements"] = [pair[0] for pair in top_pairs]
        job_data["scores"] = [pair[1] for pair in top_pairs]
    
    # Rank jobs by average similarity score (of their top achievements)
    ranked_jobs = []
    for job_key, job_data in job_groups.items():
        avg_score = sum(job_data["scores"]) / len(job_data["scores"]) if job_data["scores"] else 0
        ranked_jobs.append({
            "job_key": job_key,
            "avg_score": avg_score,
            "data": job_data
        })
    
    # Sort by average score (descending)
    ranked_jobs.sort(key=lambda x: x["avg_score"], reverse=True)
    
    # Print ranked jobs to terminal
    if ranked_jobs:
        print("\n" + "="*80)
        print("📊 RANKED JOBS BY RELEVANCE (Max 3 Achievements Per Job)")
        print("="*80)
        for idx, job_info in enumerate(ranked_jobs, 1):
            company, position, start_date, end_date = job_info["job_key"]
            avg_score = job_info["avg_score"]
            num_achievements = len(job_info["data"]["achievements"])
            print(f"\n{idx}. {position} at {company}")
            print(f"   Period: {start_date} - {end_date}")
            print(f"   Relevance Score: {avg_score:.4f}")
            print(f"   Top Achievements: {num_achievements}")
        print("\n" + "="*80 + "\n")
    
    # Select top N jobs if filtering is enabled
    if use_filtering:
        selected_jobs = ranked_jobs[:top_k_jobs]
        print(f"✅ Selected top {len(selected_jobs)} most relevant job(s) for resume generation\n")
    else:
        selected_jobs = ranked_jobs
        print(f"✅ Retrieved all {len(selected_jobs)} job(s) (full resume context)\n")  # Include all jobs for full retrieval
    
    # Build work experience from selected jobs
    for job_info in selected_jobs:
        company, position, start_date, end_date = job_info["job_key"]
        job_data = job_info["data"]
        
        resume_data['work_experience'].append({
            "position": position,
            "company": company,
            "start_date": start_date,
            "end_date": end_date,
            "location": job_data["metadata"]["location"],
            "industry": job_data["metadata"]["industry"],
            "achievements": job_data["achievements"]
        })
    
    # Sort jobs by start date (most recent first)
    resume_data['work_experience'].sort(
        key=lambda x: x['start_date'],
        reverse=True
    )
    
    # Education
    for result in education_results:
        metadata = result['metadata']
        resume_data['education']['degrees'].append({
            "degree": metadata.get('degree', ''),
            "institution": metadata.get('institution', ''),
            "dates": metadata.get('dates', '')
        })
    
    # Continuing studies
    for result in continuing_results:
        metadata = result['metadata']
        resume_data['education']['continuing_studies'].append({
            "name": metadata.get('name', ''),
            "institution": metadata.get('institution', ''),
            "completion_date": metadata.get('completion_date', '')
        })
    
    # Skills (reconstruct by category)
    for result in skills_results:
        metadata = result['metadata']
        category = metadata.get('category', 'other')
        # Parse skills from content
        content = result['content']
        if ':' in content:
            skills_text = content.split(':', 1)[1].strip()
            skills_list = [s.strip() for s in skills_text.split(',')]
            
            # Convert category to snake_case key
            category_key = category.lower().replace(' ', '_')
            resume_data['skills'][category_key] = skills_list
    
    return resume_data


def retrieve_personality_traits(job_analysis: Dict[str, Any], top_k: int = 12) -> str:
    """
    Retrieve personality traits relevant to job requirements using RAG.
    
    Uses semantic search to find personality traits that match the job's
    soft skills and keywords. Returns formatted text for cover letter inclusion.
    
    Parameters
    ----------
    job_analysis : Dict[str, Any]
        Job analysis dictionary containing at minimum:
        - soft_skills : List[str]
        - keywords : List[str]
    top_k : int, optional
        Number of personality trait chunks to retrieve (default: 12).
    
    Returns
    -------
    str
        Formatted personality traits text with section headers and descriptions.
    
    Examples
    --------
    >>> job_analysis = {
    ...     "soft_skills": ["collaboration", "problem-solving", "communication"],
    ...     "keywords": ["team", "analytical", "independent"]
    ... }
    >>> traits = retrieve_personality_traits(job_analysis)
    >>> print(traits[:50])
    'Innovative Mindset: My ability to see possibilities...'
    """
    # Initialize embeddings and vector store
    embedder = OpenAIEmbeddings()
    store = QdrantVectorStore(collection_name="personality")
    
    # Build query from job analysis
    soft_skills = job_analysis.get("soft_skills", [])
    keywords = job_analysis.get("keywords", [])
    
    query_text = f"{' '.join(soft_skills)} {' '.join(keywords)}"
    
    print(f"\n🔍 Retrieving personality traits...")
    print(f"   Query: {query_text[:100]}...")
    
    # Generate query embedding
    query_vector = embedder.embed_query(query_text)
    
    # Search personality collection
    results = store.search(query_vector, top_k=top_k)
    
    print(f"   ✅ Retrieved {len(results)} personality trait chunks")
    print(f"   Top match: {results[0]['content'][:80]}... (score: {results[0]['score']:.3f})")
    
    # Deduplicate by content to avoid repetition from overlapping chunks
    seen_starts = set()
    unique_results = []
    
    for result in results:
        # Use first 50 chars as deduplication key
        content_start = result["content"][:50]
        if content_start not in seen_starts:
            seen_starts.add(content_start)
            unique_results.append(result)
    
    # Format results
    personality_text = "\n\n".join([r["content"] for r in unique_results])
    
    return personality_text


def retrieve_portfolio_projects_hierarchical(
    job_analysis: Dict[str, Any],
    top_k_technical: int = 10,
    top_k_prompt: int = 3,
    top_k_list: int = 5
) -> Dict[str, Any]:
    """
    Retrieve portfolio projects using two-step hierarchical search.
    
    Step 1: Query technical summary chunks to get top 10 technically relevant projects.
    Step 2: Fetch full content for top 3 (for cover letter prompt) and metadata for top 5 (for list).
    
    Parameters
    ----------
    job_analysis : Dict[str, Any]
        Job analysis containing technical_skills and keywords.
    top_k_technical : int, optional
        Number of technical chunks to retrieve in first pass (default: 10).
    top_k_prompt : int, optional
        Number of full projects to include in cover letter prompt (default: 3).
    top_k_list : int, optional
        Number of projects to list at end of cover letter (default: 5).
    
    Returns
    -------
    Dict[str, Any]
        {
            "projects_for_prompt": List[Dict] with title, content, url (top 3),
            "projects_for_list": List[Dict] with title, url (top 5)
        }
    
    Examples
    --------
    >>> job_analysis = {
    ...     "technical_skills": ["Python", "Pandas", "SQL"],
    ...     "keywords": ["data pipeline", "ETL", "financial analysis"]
    ... }
    >>> result = retrieve_portfolio_projects_hierarchical(job_analysis)
    >>> len(result["projects_for_prompt"])
    3
    >>> len(result["projects_for_list"])
    5
    >>> print(result["projects_for_prompt"][0]["title"])
    'Credit Card Offer Analysis'
    """
    # Initialize embeddings and vector store
    embedder = OpenAIEmbeddings()
    store = QdrantVectorStore(collection_name="projects")
    
    # Build query from job analysis
    technical_skills = job_analysis.get("technical_skills", [])
    keywords = job_analysis.get("keywords", [])
    
    query_text = f"{' '.join(technical_skills)} {' '.join(keywords)}"
    
    print(f"\n🔍 Step 1: Retrieving technically relevant projects...")
    print(f"   Query: {query_text[:100]}...")
    
    # Generate query embedding
    query_vector = embedder.embed_query(query_text)
    
    # Step 1: Search technical summary chunks
    tech_results = store.search(
        query_vector,
        top_k=top_k_technical,
        section_filter="project_technical"
    )
    
    print(f"   ✅ Retrieved {len(tech_results)} technical summary chunks")
    if tech_results:
        print(f"   Top match: {tech_results[0]['metadata']['project_title']} (score: {tech_results[0]['score']:.3f})")
        print(f"\n   📋 Technical chunks retrieved:")
        for idx, r in enumerate(tech_results, 1):
            print(f"      {idx}. {r['metadata']['project_title']} (score: {r['score']:.3f})")
    
    # Extract unique project IDs (preserve ranking)
    seen = set()
    project_ids = []
    for r in tech_results:
        pid = r["metadata"]["project_id"]
        if pid not in seen:
            project_ids.append(pid)
            seen.add(pid)
    
    print(f"\n   🎯 Unique projects identified: {len(project_ids)}")
    
    # Step 2a: Get full content for top projects (for prompt)
    print(f"\n🔍 Step 2: Fetching full project content...")
    projects_for_prompt = []
    for idx, pid in enumerate(project_ids[:top_k_prompt], 1):
        full_result = store.search_by_metadata(
            filter_field="project_id",
            filter_value=pid,
            section_filter="project_full"
        )
        if full_result:
            projects_for_prompt.append({
                "title": full_result["metadata"]["project_title"],
                "content": full_result["content"],
                "url": full_result["metadata"]["project_url"]
            })
            print(f"   ✅ {idx}. {full_result['metadata']['project_title']}")
            print(f"      Content length: {len(full_result['content'])} chars")
    
    # Step 2b: Get metadata for top projects (for list)
    projects_for_list = []
    for pid in project_ids[:top_k_list]:
        # Fetch from either chunk type (just need metadata)
        result = store.search_by_metadata(
            filter_field="project_id",
            filter_value=pid
        )
        if result:
            projects_for_list.append({
                "title": result["metadata"]["project_title"],
                "url": result["metadata"]["project_url"],
                "tech_stack": result["metadata"].get("tech_stack", [])
            })
    
    return {
        "projects_for_prompt": projects_for_prompt,
        "projects_for_list": projects_for_list
    }


def retrieve_personality_traits(job_analysis: Dict[str, Any], top_k: int = 12) -> str:
    """
    Retrieve relevant personality traits for cover letter enhancement.
    
    Uses the "personality" collection in Qdrant with pure semantic search
    (no filtering). Handles overlapping chunks from fixed-size chunking by
    deduplicating content based on character ranges.
    
    Parameters
    ----------
    job_analysis : Dict[str, Any]
        Job analysis containing soft_skills and keywords.
    top_k : int, optional
        Number of personality chunks to retrieve (default: 12, increased to get
        more comprehensive trait coverage for cover letters).
    
    Returns
    -------
    str
        Deduplicated personality traits text.
    
    Examples
    --------
    >>> job_analysis = {
    ...     'soft_skills': ['collaboration', 'problem-solving', 'communication'],
    ...     'keywords': ['team player', 'analytical']
    ... }
    >>> traits = retrieve_personality_traits(job_analysis, top_k=8)
    >>> print(traits)
    
    Notes
    -----
    This function uses the "personality" collection with simple fixed-size
    chunking (no section awareness). Semantic search handles relevance ranking.
    Overlapping chunks (100-char overlap from 400-char chunks) are merged to
    avoid duplicate text.
    """
    
    def _merge_overlapping_chunks(results: List[Dict[str, Any]]) -> str:
        """
        Merge overlapping chunks to deduplicate content.
        
        Parameters
        ----------
        results : List[Dict[str, Any]]
            Search results with metadata containing char_start, char_end, chunk_index.
        
        Returns
        -------
        str
            Deduplicated text without section headers.
        
        Notes
        -----
        Sorts chunks by chunk_index to maintain document order,
        then removes overlapping portions based on character ranges.
        """
        if not results:
            return ""
        
        # Sort by chunk_index to maintain document order
        sorted_results = sorted(
            results,
            key=lambda x: x.get('metadata', {}).get('chunk_index', 0)
        )
        
        merged_text = []
        prev_char_end = -1
        
        for result in sorted_results:
            char_start = result.get('metadata', {}).get('char_start', 0)
            char_end = result.get('metadata', {}).get('char_end', 0)
            content = result.get('content', '')
            
            # Handle overlapping content
            if char_start < prev_char_end:
                # Calculate overlap size and skip overlapping portion
                overlap_size = prev_char_end - char_start
                if overlap_size < len(content):
                    merged_text.append(content[overlap_size:])
                    prev_char_end = char_end
            else:
                # No overlap, add full content
                merged_text.append(content)
                prev_char_end = char_end
        
        return '\n'.join(merged_text)
    
    # Initialize embeddings and vector store (use personality collection)
    embedder = OpenAIEmbeddings()
    store = QdrantVectorStore(collection_name="personality")
    
    # Construct query from soft skills and keywords
    query_parts = []
    if job_analysis.get('soft_skills'):
        query_parts.extend(job_analysis['soft_skills'])
    if job_analysis.get('keywords'):
        query_parts.extend(job_analysis['keywords'])
    
    query_text = ' '.join(query_parts) if query_parts else "personality traits"
    print(f"\n🔍 Personality Query: '{query_text}'")
    
    query_vector = embedder.embed_query(query_text)
    
    # Retrieve from personality collection without filtering
    # Semantic search handles relevance ranking
    results = store.search(
        query_vector=query_vector,
        top_k=top_k,
        section_filter=None  # No filtering - pure semantic search
    )
    
    print(f"📊 Retrieved {len(results)} chunk(s) from personality collection")
    
    # Show what was retrieved before deduplication
    for i, result in enumerate(results[:5], 1):  # Show first 5
        score = result.get('score', 0)
        preview = result.get('content', '')[:80].replace('\n', ' ')
        print(f"   {i}. Score: {score:.3f} - {preview}...")
    
    # Merge overlapping chunks and return deduplicated text
    merged = _merge_overlapping_chunks(results)
    print(f"✅ Merged and deduplicated {len(results)} chunk(s)\n")
    
    return merged

def create_resume_prompt(resume_data: Dict[str, Any], job_analysis: Dict[str, Any], job_title: str, company: str, job_description: str) -> str:
    # Here the resume_data is from the RAG retrieval
    prompt_resume = f"""
    You are an expert resume writer in the data analytics, data science, and AI engineering fields.
    Your task is to create a tailored resume that accurately reflects the candidate's qualifications and aligns with the job requirements. Output the resume in the specified JSON format only.

    First, analyze and understand the following job posting and the candidate's resume:

    Job Title
    <job_title>
    {job_title}
    </job_title>

    Company
    <company>
    {company}
    </company>

    Job Description
    <job_description>
    {job_description}
    </job_description>

    Technical Skills
    <technical_skills>
    {json.dumps(job_analysis['technical_skills'])}
    </technical_skills>

    Soft Skills
    <soft_skills>
    {json.dumps(job_analysis['soft_skills'])}
    </soft_skills>

    Keywords
    <keywords>
    {json.dumps(job_analysis['keywords'])}
    </keywords>

    Current Resume Data:
    <resume_data>
    {json.dumps(resume_data, indent=2)}
    </resume_data>

    Think through these steps internally (do not output your analysis):

    1. Professional Summary:
        - Create a concise 50-word professional summary that highlights the candidate's relevant experience and skills tailored to the job description.
       
       Constraints:
       - Use keywords from job ONLY if they appear in resume data
       - Don't use generic terms like "expert", "skilled", "seasoned", "proficient"
       - Don't claim experience with tools/technologies not in resume
       - Use exact terminology from resume when possible
       - Focus on demonstrable experience and concrete technologies
    
    2. Work Experience:
        - Select the four most relevant work experiences from resume data including the most recent. Create 3 bullet points for each experience (2 bullet points for quality engineer and quality assurance roles) using strong action verbs and quantified achievements. Ensure each bullet point is relevant to the job requirements.
    
    3. Education: 
        - List all educational qualifications from resume data. Note any specific educational requirements from job description. Ensure consistent formatting for all entries.
    
    4. Continuing Studies:
        - Select the three most relevant to the job requirements.

    CRITICAL: Output ONLY a valid JSON object with this exact structure. No text before or after:

    {{
        "professional_summary": "string",
        "work_experience": [
            {{
                "company": "string",
                "title": "string",
                "dates": "string",
                "bullet_points": ["string", "string", "string"]
            }}
        ],
        "education": [
            {{
                "degree": "string",
                "institution": "string",
                "year_completed": "string"
            }}
        ],
        "continuing_studies": [
            {{
                "name": "string",
                "institution": "string",
                "completion_date": "string"
            }}
        ]
    }}

    Don't use "Present" if the end date is before March 2026. Output pure JSON only.
    """
    return prompt_resume

# Resume Generation Data Validation
class ResumeResponse(BaseModel):
    """Model for the resume response structure."""
    professional_summary: str = ""
    work_experience: List[Dict[str, Any]] = Field(default_factory=list)
    education: List[Dict[str, Any]] = Field(default_factory=list)
    continuing_studies: List[Dict[str, Any]] = Field(default_factory=list)
    
    @model_validator(mode='before')
    @classmethod
    def extract_from_text(cls, data):
        """Extract JSON from text if the input is not already a dict."""
        if isinstance(data, dict):
            return data
            
        if isinstance(data, str):
            # Try direct JSON parsing first
            try:
                return json.loads(data)
            except json.JSONDecodeError:
                # Extract JSON using regex if direct parsing fails
                try:
                    json_pattern = r'(\{[\s\S]*\})'
                    match = re.search(json_pattern, data)
                    if match:
                        return json.loads(match.group(1))
                except:
                    pass
                
                # If all else fails, print a warning and return empty structure
                print(f"Failed to parse resume response: {data[:100]}...")
                return {
                    "professional_summary": "",
                    "work_experience": [],
                    "education": [],
                    "continuing_studies": []
                }
        
        return data
    

def claude_resume(api_key: str, prompt_resume: str) -> Dict[str, Any]:
    client = Anthropic(api_key=api_key)
    response = client.messages.create(
        model=model_name,
        max_tokens=2048,
        temperature=0.6,
        messages=[{"role": "user", "content": prompt_resume}]
    )
    # Use Pydantic model to validate and extract the response
    response_text = response.content[0].text
    try:
        resume_data = ResumeResponse.model_validate(response_text)
        return resume_data.model_dump()
    except Exception as e:
        print(f"Error parsing resume response with Pydantic: {e}")
        # Fallback to simpler parsing if Pydantic validation fails
        try:
            return json.loads(response_text)
        except:
            print(f"JSON parsing failed. Raw response: {response_text[:150]}...")
            # Last resort: return empty data structure
            return {
                "professional_summary": "",
                "work_experience": [],
                "education": [],
                "continuing_studies": []
            }

print("Resume generation functions defined!")

# %%
# Cover Letter Generation Functions

def create_cover_letter_prompt(resume_data: Dict[str, Any], job_analysis: Dict[str, Any], job_title: str, company: str, personality_traits: str = "", portfolio_projects: Dict[str, Any] = None) -> str:
    """
    Generate cover letter prompt with portfolio project integration.
    
    Parameters
    ----------
    resume_data : Dict[str, Any]
        Generated resume content.
    job_analysis : Dict[str, Any]
        Job analysis with technical_skills, soft_skills, keywords.
    job_title : str
        Job title.
    company : str
        Company name.
    job_description : str
        Full job description text.
    personality_traits : str, optional
        Relevant personality traits text.
    portfolio_projects : Dict[str, Any], optional
        Portfolio projects from get_portfolio_projects tool with keys:
        - projects_for_prompt: List[Dict] (top 3 with full content)
        - projects_for_list: List[Dict] (top 5 with title+URL)
    
    Returns
    -------
    str
        Cover letter generation prompt.
    """
    # Format portfolio projects for prompt
    projects_context = ""
    
    if portfolio_projects:
        # Extract projects for narrative
        if portfolio_projects.get('projects_for_prompt'):
            projects_context = "\n\n".join([
                f"### {p['title']}\n{p['content']}" 
                for p in portfolio_projects['projects_for_prompt']
            ])
    
    prompt_cover_letter = f"""
    You are an expert cover letter writer in the data analytics, data science, and AI engineering fields.
    Your task is to create a tailored cover letter that accurately reflects the candidate's qualifications and aligns with the job requirements. Output the cover letter in the specified JSON format only. 

    Here is a summary of the job posting:
    <job_summary>
    {job_analysis.get('job_summary', 'No job summary provided.')}
    </job_summary>

    Here is the candidate's resume data:
    <resume_data>
    {json.dumps(resume_data, indent=2)}
    </resume_data>

    Here is the job title:
    <job_title>
    {job_title}
    </job_title>

    Here is the company name:
    <company>
    {company}
    </company>

    Here are the candidate's relevant personality traits:
    <personality_traits>
    {personality_traits if personality_traits else "No specific personality traits provided."}
    </personality_traits>

    Here are the candidate's relevant portfolio projects:
    <portfolio_projects_context>
    {projects_context if projects_context else "No portfolio projects provided."}
    </portfolio_projects_context>

    Here are the relevant soft skills:
    <soft_skills>
    {json.dumps(job_analysis['soft_skills'])}
    </soft_skills>

    Here are the important keywords:
    <keywords>
    {json.dumps(job_analysis['keywords'])}
    </keywords>

    Structure the cover letter in 3 short paragraphs (opening, body, closing), under 250 words total:
        - Opening: Demonstrate why you are interested in the role and company and how your skills solve industry challenges or specific problems mentioned in the job summary.
        - Body: Highlight soft skills matching job requirements, naturally include in relevant personality traits and reference 1-2 portfolio projects to demonstrate skill.
        - Closing: Reiterate enthusiasm for the role, summarize key qualifications, and include a call to action for next steps.
    
    Use a casual friendly tone, avoiding formal language or clichés.
    Use present tense (except for past experiences).
    Avoid words like "expert", "skilled", "seasoned", "excited".
    Focus on highlighting soft skills that match the job requirements, weaving in relevant personality traits and referencing 1-2 portfolio projects to demonstrate fit.
    Use keywords naturally throughout.
    Use the provided personality traits to strengthen the narrative and show cultural/role fit.
    Only use information from the resume that matches the job requirements.
 

    CRITICAL: Output ONLY a valid JSON object with this exact structure. No text before or after:

    {{
        "opening_paragraph": "string (single paragraph)",
        "body_paragraph": "string (single paragraph, mention 1-2 portfolio projects if provided)",
        "closing_paragraph": "string (single paragraph)"
    }}

    """
    return prompt_cover_letter

# Cover Letter Generation Data Validation
class CoverLetterResponse(BaseModel):
    """Model for the cover letter response structure."""
    opening_paragraph: str = "Unable to generate opening paragraph."
    body_paragraph: str = "Unable to generate body paragraph."
    closing_paragraph: str = "Unable to generate closing paragraph."
    
    @model_validator(mode='before')
    @classmethod
    def extract_from_text(cls, data):
        """Extract JSON from text if the input is not already a dict."""
        if isinstance(data, dict):
            return data
            
        if isinstance(data, str):
            # Try direct JSON parsing first
            try:
                return json.loads(data)
            except json.JSONDecodeError:
                # Extract JSON using regex if direct parsing fails
                try:
                    json_pattern = r'(\{[\s\S]*\})'
                    match = re.search(json_pattern, data)
                    if match:
                        return json.loads(match.group(1))
                except:
                    pass
                
                # If all else fails, return default content
                return {
                    "opening_paragraph": "Unable to generate opening paragraph.",
                    "body_paragraph": "Unable to generate body paragraph.",
                    "closing_paragraph": "Unable to generate closing paragraph."
                }
        
        return data

def claude_cover_letter(api_key: str, prompt_cover_letter: str) -> Dict[str, Any]:
    client = Anthropic(api_key=api_key)
    response = client.messages.create(
        model=model_name,
        max_tokens=4096,
        temperature=0.7,
        messages=[{"role": "user", "content": prompt_cover_letter}]
    )
    # Use Pydantic model to validate and extract the response
    response_text = response.content[0].text
    try:
        cover_letter_data = CoverLetterResponse.model_validate(response_text)
        return cover_letter_data.model_dump()
    except Exception as e:
        print(f"Error parsing response with Pydantic: {e}")
        # Fallback to simpler parsing if Pydantic validation fails
        try:
            return json.loads(response_text)
        except:
            print(f"JSON parsing failed. Raw response: {response_text[:150]}...")
            # Last resort: return empty data structure
            return {
                "opening_paragraph": "Unable to generate opening paragraph.",
                "body_paragraph": "Unable to generate body paragraph.",
                "closing_paragraph": "Unable to generate closing paragraph."
            }

print("Cover letter generation functions defined!")

# %%
# Document Creation Functions

def _setup_document() -> Document:
    """
    Initialize and configure a Word document with standard settings.
    
    This helper function creates a new document and applies global formatting
    settings including margins, fonts, footer with page numbers, and header spacing.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    Document
        A configured python-docx Document object with standard formatting applied.
    
    Notes
    -----
    - Sets Helvetica font as default
    - Margins: Top=0.3", Bottom=0.2", Left/Right=0.5"
    - Adds right-aligned page number in footer
    - Footer distance: 0.3", Header distance: 0.3"
    
    Examples
    --------
    >>> doc = _setup_document()
    >>> # Document is ready for content addition
    """
    doc = Document()
    
    # Add footer with page number
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = paragraph.add_run()
    
    # Begin field char
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    
    # Field instruction text
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._element.append(instrText)
    
    # End field char
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)

    # Set footer margins
    section.footer_distance = Inches(0.3)
    
    # Set header margins
    section.header_distance = Inches(0.3)

    # GLOBAL DOCUMENT SETTINGS
    style = doc.styles['Normal']
    style.font.name = 'Helvetica'

    # Set margin sizes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.2)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    return doc


def _add_resume_section(doc: Document, resume: Dict[str, Any], resume_ale: Dict[str, Any]) -> None:
    """
    Add resume content section to a Word document.
    
    This helper function adds a complete resume section including header with contact
    information, professional summary, work experience, continuing studies, and education.
    
    Parameters
    ----------
    doc : Document
        The python-docx Document object to add content to.
    resume : Dict[str, Any]
        Generated resume content dictionary with keys:
        - professional_summary: str
        - work_experience: List[Dict] with company, title, dates, bullet_points
        - continuing_studies: List[Dict] with name, institution, completion_date
        - education: List[Dict] with degree, institution, year_completed
    resume_ale : Dict[str, Any]
        Candidate's base resume data from YAML file containing personal_information.
    
    Returns
    -------
    None
        Modifies the document in-place.
    
    Notes
    -----
    - Header includes name (22pt, bold, centered) and contact info (10pt, blue)
    - Professional summary is justified text
    - Work experience uses tab stops for right-aligned dates
    - Bullet points are indented 0.15 inches
    - Section headers are centered and bold
    - Phone number: 778.223.8536
    
    Examples
    --------
    >>> doc = _setup_document()
    >>> _add_resume_section(doc, resume_data, base_data)
    """
    # Name
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(resume_ale['personal_information']['full_name'])
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title.paragraph_format.space_after = Pt(1.15)
    
    # Contact info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    email_run = contact.add_run(resume_ale['personal_information']['email'])
    email_run.font.size = Pt(10)
    email_run.font.color.rgb = RGBColor(23, 54, 93)
    separator_run = contact.add_run(' | ')
    linkedin_run = contact.add_run(resume_ale['personal_information']['linkedin'])
    linkedin_run.font.size = Pt(10)
    linkedin_run.font.color.rgb = RGBColor(23, 54, 93)
    separator_run = contact.add_run(' | ')
    github_run = contact.add_run(resume_ale['personal_information']['github'])
    github_run.font.size = Pt(10)
    github_run.font.color.rgb = RGBColor(23, 54, 93)
    separator_run = contact.add_run(' | ')
    phone_run = contact.add_run("778.223.8536")
    phone_run.font.size = Pt(10)
    phone_run.font.color.rgb = RGBColor(23, 54, 93)
    
    # Professional summary
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary.add_run(resume['professional_summary'])
    summary.paragraph_format.space_after = Pt(10)
    
    # Experience
    exp_header = doc.add_paragraph()
    exp_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    exp_header.add_run("EXPERIENCE").bold = True
    exp_header.paragraph_format.space_after = Pt(0)
    
    # Set up the tab stop for right-aligned dates
    tab_stops = doc.styles['Normal'].paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.6), WD_TAB_ALIGNMENT.RIGHT)
    
    # Add work experience with right-aligned dates
    for job in resume['work_experience']:
        # Create paragraph for job header
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        
        # Add job title and company
        title_run = p.add_run(f"{job['title']} | ")
        title_run.bold = True
        company_run = p.add_run(job['company'])
        company_run.italic = True
        
        # Add right-aligned date using tab
        date_run = p.add_run('\t' + job['dates'])
        date_run.italic = True
        date_run.bold = True
        
        # Add bullet points
        for point in job['bullet_points']:
            bullet = doc.add_paragraph()
            bullet.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            bullet.paragraph_format.left_indent = Inches(0.15)
            bullet.add_run('• ' + point)
            bullet.paragraph_format.space_after = Pt(6)

    # Add Continuing Studies section
    cont_header = doc.add_paragraph()
    cont_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cont_header.add_run("CONTINUING STUDIES").bold = True
    cont_header.paragraph_format.space_before = Pt(10)
    cont_header.paragraph_format.space_after = Pt(0)

    # Add continuing studies entries
    for cont in resume['continuing_studies']:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        name_run = p.add_run(f"{cont['name']} | ")
        name_run.bold = True
        inst_run = p.add_run(cont['institution'])
        inst_run.italic = True
        
        # Add completion date (right-aligned)
        completion_date = p.add_run('\t' + cont['completion_date'])
        completion_date.italic = True
        completion_date.bold = True
    
    # Add Education section
    edu_header = doc.add_paragraph()
    edu_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    edu_header.add_run("EDUCATION").bold = True
    edu_header.paragraph_format.space_before = Pt(10)
    edu_header.paragraph_format.space_after = Pt(0)
    
    # Add education entries
    for edu in resume['education']:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        degree_run = p.add_run(f"{edu['degree']} | ")
        degree_run.bold = True
        inst_run = p.add_run(edu['institution'])
        inst_run.italic = True
        
        # Add year (right-aligned)
        completion_date = p.add_run('\t' + edu['year_completed'])
        completion_date.italic = True
        completion_date.bold = True


def _add_cover_letter_section(doc: Document, cover_letter: Dict[str, Any], resume_ale: Dict[str, Any], company: str, portfolio_projects: Optional[Dict[str, Any]] = None, hiring_manager_greeting: str = "Dear Hiring Manager:") -> None:
    """
    Add cover letter content section to a Word document.
    
    This helper function adds a complete cover letter section including header with
    contact information, date, company name, greeting, three content paragraphs,
    and portfolio projects list.
    
    Parameters
    ----------
    doc : Document
        The python-docx Document object to add content to.
    cover_letter : Dict[str, Any]
        Generated cover letter content dictionary with keys:
        - opening_paragraph: str
        - body_paragraph: str
        - closing_paragraph: str
    resume_ale : Dict[str, Any]
        Candidate's base resume data from YAML file containing personal_information.
    company : str
        Company name for the cover letter recipient.
    portfolio_projects : Optional[Dict[str, Any]], optional
        Portfolio projects from vector store with keys:
        - projects_for_list: List[Dict] with title and url
    
    Returns
    -------
    None
        Modifies the document in-place.
    
    Notes
    -----
    - Header includes name (22pt, bold, centered) and contact info (10pt, blue)
    - Date is right-aligned in "Month DD, YYYY" format
    - Greeting uses "Dear Hiring Manager:"
    - Three paragraphs: opening, body, closing (12pt spacing between)
    - Signature includes "Sincerely," and candidate name
    - Portfolio projects list appears after signature if projects_for_list exists
    - Phone number: 778.223.8536
    
    Examples
    --------
    >>> doc = _setup_document()
    >>> _add_cover_letter_section(doc, cover_letter_data, base_data, "Tech Corp", portfolio_projects)
    """
    # Name
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(resume_ale['personal_information']['full_name'])
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title.paragraph_format.space_after = Pt(1.15)
    
    # Contact info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    email_run = contact.add_run(resume_ale['personal_information']['email'])
    email_run.font.size = Pt(10)
    email_run.font.color.rgb = RGBColor(23, 54, 93)
    separator_run = contact.add_run(' | ')
    linkedin_run = contact.add_run(resume_ale['personal_information']['linkedin'])
    linkedin_run.font.size = Pt(10)
    linkedin_run.font.color.rgb = RGBColor(23, 54, 93)
    separator_run = contact.add_run(' | ')
    github_run = contact.add_run(resume_ale['personal_information']['github'])
    github_run.font.size = Pt(10)
    github_run.font.color.rgb = RGBColor(23, 54, 93)
    separator_run = contact.add_run(' | ')
    phone_run = contact.add_run("778.223.8536")
    phone_run.font.size = Pt(10)
    phone_run.font.color.rgb = RGBColor(23, 54, 93)

    # Current Date
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date.add_run(datetime.now().strftime('%B %d, %Y'))

    # Add space after the date
    date.paragraph_format.space_after = Pt(24)
    
    # Recipient info (Company name)    
    recipient = doc.add_paragraph()
    recipient.alignment = WD_ALIGN_PARAGRAPH.LEFT
    recipient.add_run(company)
    recipient.paragraph_format.space_after = Pt(24)
    
    # Greeting
    greeting = doc.add_paragraph()
    greeting.alignment = WD_ALIGN_PARAGRAPH.LEFT
    greeting.add_run(hiring_manager_greeting)
    greeting.paragraph_format.space_after = Pt(12)
    
    # Opening paragraph
    opening = doc.add_paragraph()
    opening.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    opening.add_run(cover_letter['opening_paragraph'])
    opening.paragraph_format.space_after = Pt(12)
    
    # Body paragraph
    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.add_run(cover_letter['body_paragraph'])
    body.paragraph_format.space_after = Pt(12)
    
    # Closing paragraph
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    closing.add_run(cover_letter['closing_paragraph'])
    closing.paragraph_format.space_after = Pt(24)
    
    # Signature
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.LEFT
    signature.add_run('Sincerely,\n\nAlejandro Leiva')
    signature.paragraph_format.space_after = Pt(24)
    
    # Portfolio projects list (if available)
    if portfolio_projects and portfolio_projects.get('projects_for_list') and len(portfolio_projects['projects_for_list']) > 0:
        projects_header = doc.add_paragraph()
        projects_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        projects_header.add_run("Here is a list of relevant portfolio projects:")
        projects_header.paragraph_format.space_after = Pt(6)
        
        for project in portfolio_projects['projects_for_list']:
            # Project title and URL
            project_item = doc.add_paragraph()
            project_item.alignment = WD_ALIGN_PARAGRAPH.LEFT
            project_item.add_run(f"- {project['title']}: {project['url']}")
            project_item.paragraph_format.space_after = Pt(2)
            project_item.paragraph_format.left_indent = Inches(0.25)
            # Add tech stack
            tech_stack_text = ", ".join(project['tech_stack'])
            project_item = doc.add_paragraph()
            project_item.alignment = WD_ALIGN_PARAGRAPH.LEFT
            tech_run = project_item.add_run(f"Tech Stack: {tech_stack_text}")
            project_item.paragraph_format.space_after = Pt(6)
            project_item.paragraph_format.left_indent = Inches(0.25)


def create_resume_document(resume: Dict[str, Any], resume_ale: Dict[str, Any]) -> Document:
    """
    Create a Word document containing only a resume.
    
    This function generates a single-page Word document with the candidate's resume,
    including professional summary, work experience, continuing studies, and education.
    
    Parameters
    ----------
    resume : Dict[str, Any]
        Generated resume content dictionary with keys:
        - professional_summary: str
        - work_experience: List[Dict] with company, title, dates, bullet_points
        - continuing_studies: List[Dict] with name, institution, completion_date
        - education: List[Dict] with degree, institution, year_completed
    resume_ale : Dict[str, Any]
        Candidate's base resume data from YAML file containing personal_information.
    
    Returns
    -------
    Document
        A python-docx Document object containing the formatted resume.
    
    External Files Required
    -----------------------
    None. All data is passed as parameters.
    
    Examples
    --------
    Create a resume document:
    
    >>> doc = create_resume_document(resume_data, base_data)
    >>> doc.save('resume.docx')
    """
    doc = _setup_document()
    _add_resume_section(doc, resume, resume_ale)
    return doc


def create_cover_letter_document(cover_letter: Dict[str, Any], resume_ale: Dict[str, Any], company: str, portfolio_projects: Optional[Dict[str, Any]] = None, hiring_manager_greeting: str = "Dear Hiring Manager:") -> Document:
    """
    Create a Word document containing only a cover letter.
    
    This function generates a single-page Word document with a tailored cover letter,
    including header, date, company address, greeting, and content paragraphs.
    
    Parameters
    ----------
    cover_letter : Dict[str, Any]
        Generated cover letter content dictionary with keys:
        - opening_paragraph: str
        - body_paragraph: str
        - closing_paragraph: str
    resume_ale : Dict[str, Any]
        Candidate's base resume data from YAML file containing personal_information.
    company : str
        Company name for the cover letter recipient.
    portfolio_projects : Optional[Dict[str, Any]], optional
        Portfolio projects from vector store with projects_for_list.
    
    Returns
    -------
    Document
        A python-docx Document object containing the formatted cover letter.
    
    External Files Required
    -----------------------
    None. All data is passed as parameters.
    
    Examples
    --------
    Create a cover letter document:
    
    >>> doc = create_cover_letter_document(cover_letter_data, base_data, "Tech Corp", portfolio_projects)
    >>> doc.save('cover_letter.docx')
    """
    doc = _setup_document()
    _add_cover_letter_section(doc, cover_letter, resume_ale, company, portfolio_projects, hiring_manager_greeting)
    return doc


def create_resume_coverletter(resume: Dict[str, Any], resume_ale: Dict[str, Any], cover_letter: Dict[str, Any], company: str, portfolio_projects: Optional[Dict[str, Any]] = None, hiring_manager_greeting: str = "Dear Hiring Manager:") -> Document:
    """
    Create a Word document containing both resume and cover letter.
    
    This function generates a two-page Word document with the candidate's resume
    on the first page and a tailored cover letter on the second page.
    
    Parameters
    ----------
    resume : Dict[str, Any]
        Generated resume content dictionary with keys:
        - professional_summary: str
        - work_experience: List[Dict] with company, title, dates, bullet_points
        - continuing_studies: List[Dict] with name, institution, completion_date
        - education: List[Dict] with degree, institution, year_completed
    resume_ale : Dict[str, Any]
        Candidate's base resume data from YAML file containing personal_information.
    cover_letter : Dict[str, Any]
        Generated cover letter content dictionary with keys:
        - opening_paragraph: str
        - body_paragraph: str
        - closing_paragraph: str
    company : str
        Company name for the cover letter recipient.
    portfolio_projects : Optional[Dict[str, Any]], optional
        Portfolio projects from vector store with projects_for_list.
    
    Returns
    -------
    Document
        A python-docx Document object containing both the resume and cover letter.
    
    External Files Required
    -----------------------
    None. All data is passed as parameters.
    
    Notes
    -----
    This function maintains backward compatibility with existing code that expects
    a combined resume and cover letter document.
    
    Examples
    --------
    Create a complete application package:
    
    >>> doc = create_resume_coverletter(resume_data, base_data, cover_letter_data, "Tech Corp", portfolio_projects)
    >>> doc.save('application.docx')
    """
    # Setup document with standard formatting
    doc = _setup_document()
    
    # Add resume section
    _add_resume_section(doc, resume, resume_ale)
    
    # Add page break before cover letter
    if doc.paragraphs:
        last_para = doc.paragraphs[-1]
        last_para.add_run().add_break(WD_BREAK.PAGE)
    else:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    # Add cover letter section
    _add_cover_letter_section(doc, cover_letter, resume_ale, company, portfolio_projects, hiring_manager_greeting)
    
    return doc

def convert_word_to_pdf(file_path: str) -> str:
    pdf_path = file_path.replace('.docx', '.pdf')
    pythoncom.CoInitialize()  # Initialize COM
    convert(file_path, pdf_path)
    pythoncom.CoUninitialize()  # Uninitialize COM
    return pdf_path

print("Document creation functions defined!")

# %%
# # Orchestrating Function
# def create_resume(job_url: str, brightdata_api_key: str, resume_yaml_path: str = 'resume_ale.yaml') -> tuple[str, str, Dict[str, Any]]:
#     """
#     Orchestrate the complete resume and cover letter generation process.
    
#     This function serves as the main entry point for generating tailored resumes
#     and cover letters from job postings. It extracts job information, analyzes
#     requirements, generates customized content using AI, and creates professional
#     documents.
    
#     Parameters
#     ----------
#     job_url : str
#         URL of the job posting (LinkedIn or Indeed supported).
#     brightdata_api_key : str
#         BrightData API key for job extraction.
#     resume_yaml_path : str, optional
#         Path to the YAML file containing resume data. Default is 'resume_ale.yaml'.
    
#     Returns
#     -------
#     tuple[str, str, Dict[str, Any]]
#         A tuple containing the paths to the generated Word document and PDF, and the job dictionary.
#         Format: (word_document_path, pdf_document_path, job_data_dict)
    
#     Raises
#     ------
#     ValueError
#         If required environment variables are missing or job extraction fails.
#     FileNotFoundError
#         If the resume YAML file is not found.
#     Exception
#         For any other errors during the process.
    
#     Examples
#     --------
#     Generate documents for a LinkedIn job:
    
#     >>> doc_path, pdf_path = main("https://www.linkedin.com/jobs/view/12345")
#     >>> print(f"Generated: {doc_path}, {pdf_path}")
#     Generated: Data_Scientist_Company_2025_09_02.docx, Data_Scientist_Company_2025_09_02.pdf
    
#     Use a custom resume file:
    
#     >>> doc_path, pdf_path = main("https://ca.indeed.com/viewjob?jk=67890", "my_resume.yaml")
#     """
#     try:
#         # Step 1: Extract job information
#         print("Extracting job information from URL")
#         job_dict, job_df = extract_job(job_url, brightdata_api_key)
#         job_title = job_dict.get('job_title', 'Unknown Title')
#         company = job_dict.get('company_name', 'Unknown Company')
#         job_description = job_dict.get('description_text', 'No description text available')
#         print(f"Extracted: {job_title} at {company}")
        
#         # Step 2: Load resume data
#         print("Loading resume data")
#         resume_data = load_resume_yaml(resume_yaml_path)
        
#         # Step 3: Analyze job description
#         print("Analyzing job description")
#         prompt_analysis = create_analysis_prompt(job_title, company, job_description)
#         if not key:
#             raise ValueError("CLAUDE_API_KEY environment variable is required.")
#         job_analysis = claude_analysis(key, prompt_analysis)
        
#         # Step 4: Generate tailored resume
#         print("Generating tailored resume")
#         prompt_resume = create_resume_prompt(resume_data, job_analysis, job_title, company, job_description)
#         resume = claude_resume(key, prompt_resume)
        
#         # Step 5: Generate tailored cover letter
#         print("Generating tailored cover letter")
#         prompt_cover_letter = create_cover_letter_prompt(resume, job_analysis, job_title, company, job_description)
#         cover_letter = claude_cover_letter(key, prompt_cover_letter)
        
#         # Step 6: Create and save documents
#         print("Creating and saving documents")
#         today = datetime.today().strftime('%Y_%m_%d_%H_%M')
#         doc = create_resume_coverletter(resume, resume_data, cover_letter, company)
#         doc_path = f'Alejandro_Leiva_{job_title}_{company}_{today}.docx'
#         doc.save(doc_path)
#         pdf_path = convert_word_to_pdf(doc_path)
        
#         return doc_path, pdf_path, job_dict
        
#     except Exception as e:
#         print(f"❌ Error in main process: {e}")
#         raise


