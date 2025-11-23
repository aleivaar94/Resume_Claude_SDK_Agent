# %%
# Import Required Libraries
from anthropic import Anthropic
import json
import yaml
import re
from datetime import datetime
import os
from dotenv import load_dotenv
import textwrap

# Data validation & handling
from typing import List, Dict, Any, Optional
from pydantic import BaseModel, Field, model_validator

# Word Document
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# PDF
from docx2pdf import convert
import pythoncom

# BrightData script (for job scraping)
from brightdata import extract_job

# %%
# Load Environment Variables
load_dotenv()

key = os.getenv("CLAUDE_API_KEY")
if not key:
    raise ValueError("CLAUDE_API_KEY environment variable is required.")

model_name = "claude-sonnet-4-5"

# %%
# Job Analysis Functions
def create_analysis_prompt(job_title: str, company: str, job_description: str, language: str = "English") -> str:
    language_instruction = ""
    if language == "Spanish":
        language_instruction = """
        
        IMPORTANT: Generate your analysis in Spanish. All skill names, keywords, and categories should be in Spanish.
        """
    
    prompt_analysis = f"""
    You are an experienced hiring manager tasked with analyzing job descriptions to extract key information. Your goal is to provide accurate and relevant data that can be used in the recruitment process.{language_instruction}

    Here is the job information you need to analyze:

    <job_description>
    {job_description}
    </job_description>

    <job_title>
    {job_title}
    </job_title>

    <company>
    {company}
    </company>

    Your task is to extract the following information from the job details provided:

    1. technical_skills: An array of technical skills required or preferred for the role. Look for specific technologies, programming languages, tools, or technical methodologies mentioned.
    2. soft_skills: An array of soft skills or personal attributes required or preferred for the role. Identify personal attributes, interpersonal skills, or work style preferences described.
    3. keywords: An array of keywords relevant to the role, including industry-specific terms, tools, or methodologies. Extract key terms that are frequently mentioned or seem particularly important to the role or industry.

    Output the information into a JSON object only with the following structure:

    - "technical_skills": ["skill1", "skill2", "skill3"],
    - "soft_skills": ["skill1", "skill2", "skill3"],
    - "keywords": ["keyword1", "keyword2", "keyword3"]

    Each array should contain at least 3 items if possible and a maximum of 10, but don't include irrelevant information just to meet this number.

    Remember to focus on accuracy and relevance in your analysis and output.
    """
    return prompt_analysis

# Job Analysis Response Data Validation
class JobAnalysisResponse(BaseModel):
    """Model for the job analysis response structure."""
    technical_skills: List[str] = Field(default_factory=list)
    soft_skills: List[str] = Field(default_factory=list)
    keywords: List[str] = Field(default_factory=list)
    
    @model_validator(mode='before')
    @classmethod
    def extract_from_text(cls, data):
        """Extract JSON from text if the input is not already a dict."""
        if isinstance(data, dict):
            return data
            
        if isinstance(data, str):
            # Try direct JSON parsing first
            try:
                return json.loads(data)
            except json.JSONDecodeError:
                # Extract JSON using regex if direct parsing fails
                try:
                    json_pattern = r'(\{[\s\S]*\})'
                    match = re.search(json_pattern, data)
                    if match:
                        return json.loads(match.group(1))
                except:
                    pass
                
                # If we can't extract valid JSON, try to create a structured response
                # by looking for keyword patterns
                skills_pattern = {
                    "technical_skills": r'"?technical_skills"?\s*:\s*\[(.*?)\]',
                    "soft_skills": r'"?soft_skills"?\s*:\s*\[(.*?)\]',
                    "keywords": r'"?keywords"?\s*:\s*\[(.*?)\]'
                }
                
                extracted_data = {}
                for key, pattern in skills_pattern.items():
                    match = re.search(pattern, data, re.DOTALL)
                    if match:
                        # Extract the items in the array
                        items_text = match.group(1)
                        # Parse the items, handling quoted strings
                        items = []
                        for item in re.findall(r'"([^"]*)"', items_text):
                            items.append(item)
                        if items:
                            extracted_data[key] = items
                
                if extracted_data:
                    return extracted_data
                    
                # If all else fails, print a warning and return empty lists
                print(f"Failed to parse analysis response: {data[:100]}...")
                return {
                    "technical_skills": [],
                    "soft_skills": [],
                    "keywords": []
                }
        
        return data
    

def claude_analysis(api_key: str, prompt_analysis: str) -> Dict[str, Any]:
    client = Anthropic(api_key=api_key)
    response = client.messages.create(
        model=model_name,
        max_tokens=2048,
        temperature=0.5,
        messages=[{"role": "user", "content": prompt_analysis}]
    )
    # Use Pydantic model to validate and extract the response
    response_text = response.content[0].text
    try:
        analysis_data = JobAnalysisResponse.model_validate(response_text)
        return analysis_data.model_dump()
    except Exception as e:
        print(f"Error parsing response with Pydantic: {e}")
        # Fallback to simpler parsing if Pydantic validation fails
        try:
            return json.loads(response_text)
        except:
            print(f"JSON parsing failed. Raw response: {response_text[:150]}...")
            # Last resort: return empty data structure
            return {
                "technical_skills": [],
                "soft_skills": [],
                "keywords": []
            }

def format_claude_analysis_response(skills_dict):
    formatted_text = ""    
    for category, items in skills_dict.items():
        # Convert category from snake_case to Title Case
        category_title = ' '.join(word.capitalize() for word in category.split('_'))
        # Section header
        formatted_text += f"## {category_title}\n"
        # Add bullet points for each item
        for item in items:
            formatted_text += f"• {item}\n"
        # Add extra line between sections
        formatted_text += "\n"
    return formatted_text

# %%
# Resume Generation Functions

def load_resume_yaml(file_path: str) -> Dict[Any, Any]:
    with open(file_path, 'r') as file:
        resume_data = yaml.safe_load(file)
    return resume_data

def create_resume_prompt(resume_data: Dict[str, Any], job_analysis: Dict[str, Any], job_title: str, company: str, job_description: str, language: str = "English", country: str = "Canada") -> str:
    language_instruction = ""
    if language == "Spanish":
        language_instruction = """
        
        CRITICAL: Generate the entire resume in Spanish. This includes:
        - Professional summary in Spanish
        - All job titles, company names, and bullet points in Spanish
        - Education section in Spanish
        - All field names and content must be in Spanish
        """
    
    # Add Mexico-specific instructions for English fluency
    mexico_instruction = ""
    if country == "Mexico":
        mexico_instruction = """
        
        In the professional summary, explicitly include that you speak fluent English. This is crucial for the Mexican job market where English proficiency is highly valued by employers.
        """
    
    prompt_resume = f"""
    You are an expert resume writer tasked with creating a highly targeted, achievement-based resume that aligns precisely with given job requirements while accurately representing a candidate's experience.{language_instruction}{mexico_instruction}

    First, review the following information carefully:

    Job Title
    <job_title>
    {job_title}
    </job_title>

    Company
    <company>
    {company}
    </company>

    Job Description
    <job_description>
    {job_description}
    </job_description>

    Technical Skills
    <technical_skills>
    {json.dumps(job_analysis['technical_skills'])}
    </technical_skills>

    Soft Skills
    <soft_skills>
    {json.dumps(job_analysis['soft_skills'])}
    </soft_skills>

    Keywords
    <keywords>
    {json.dumps(job_analysis['keywords'])}
    </keywords>

    Current Resume Data:
    <resume_data>
    {json.dumps(resume_data, indent=2)}
    </resume_data>

    Think through these steps internally (do not output your analysis):

    1. Professional Summary: 
       CRITICAL - You must ONLY use terminology, skills, and technologies that explicitly appear in the resume data.
       
       Step 1a: Extract ONLY the technologies, tools, and skills explicitly mentioned in the candidate's work experience and continuing studies.
       Step 1b: From the job description, identify which required skills/technologies match those found in Step 1a.
       Step 1c: Calculate years of experience ONLY from the resume's work_experience dates for the matching role/field.
       Step 1d: Draft a 50-word summary using ONLY the intersection from Step 1b.
       
       Constraints:
       - Do NOT introduce terminology not present in resume data
       - Do NOT paraphrase resume skills to match job requirements
       - Do NOT use generic terms like "expert", "skilled", "seasoned", "proficient"
       - Do NOT claim experience with tools/technologies not in resume
       - DO use exact terminology from resume when possible
       - DO focus on demonstrable experience and concrete technologies
       - DO incorporate keywords from job ONLY if they appear in resume data
       {f"- DO include 'fluent English' or 'English proficiency'" if country == "Mexico" else ""}
       
       If a job requirement has no match in resume, omit it from summary.
    
    2. Work Experience: Select the four most relevant work experiences from resume data including the most recent. Create 3 bullet points for each experience (2 bullet points for quality engineer and quality assurance roles) using strong action verbs and quantified achievements. Ensure each bullet point is relevant to the job requirements.
    
    3. Education: List all educational qualifications from resume data. Note any specific educational requirements from job description. Ensure consistent formatting for all entries.
    
    4. Continuing Studies: List all certifications and continuing education courses from resume data. Identify the four most relevant to the job requirements.

    CRITICAL: Output ONLY a valid JSON object with this exact structure. No text before or after:

    {{
        "professional_summary": "string",
        "work_experience": [
            {{
                "company": "string",
                "title": "string",
                "dates": "string",
                "bullet_points": ["string", "string", "string"]
            }}
        ],
        "education": [
            {{
                "degree": "string",
                "institution": "string",
                "year_completed": "string"
            }}
        ],
        "continuing_studies": [
            {{
                "name": "string",
                "institution": "string",
                "completion_date": "string"
            }}
        ]
    }}

    Don't use "Present" if the end date is before April 2025. Output pure JSON only.
    """
    return prompt_resume

# Resume Generation Data Validation
class ResumeResponse(BaseModel):
    """Model for the resume response structure."""
    professional_summary: str = ""
    work_experience: List[Dict[str, Any]] = Field(default_factory=list)
    education: List[Dict[str, Any]] = Field(default_factory=list)
    continuing_studies: List[Dict[str, Any]] = Field(default_factory=list)
    
    @model_validator(mode='before')
    @classmethod
    def extract_from_text(cls, data):
        """Extract JSON from text if the input is not already a dict."""
        if isinstance(data, dict):
            return data
            
        if isinstance(data, str):
            # Try direct JSON parsing first
            try:
                return json.loads(data)
            except json.JSONDecodeError:
                # Extract JSON using regex if direct parsing fails
                try:
                    json_pattern = r'(\{[\s\S]*\})'
                    match = re.search(json_pattern, data)
                    if match:
                        return json.loads(match.group(1))
                except:
                    pass
                
                # If all else fails, print a warning and return empty structure
                print(f"Failed to parse resume response: {data[:100]}...")
                return {
                    "professional_summary": "",
                    "work_experience": [],
                    "education": [],
                    "continuing_studies": []
                }
        
        return data
    

def claude_resume(api_key: str, prompt_resume: str) -> Dict[str, Any]:
    client = Anthropic(api_key=api_key)
    response = client.messages.create(
        model=model_name,
        max_tokens=2048,
        temperature=0.6,
        messages=[{"role": "user", "content": prompt_resume}]
    )
    # Use Pydantic model to validate and extract the response
    response_text = response.content[0].text
    try:
        resume_data = ResumeResponse.model_validate(response_text)
        return resume_data.model_dump()
    except Exception as e:
        print(f"Error parsing resume response with Pydantic: {e}")
        # Fallback to simpler parsing if Pydantic validation fails
        try:
            return json.loads(response_text)
        except:
            print(f"JSON parsing failed. Raw response: {response_text[:150]}...")
            # Last resort: return empty data structure
            return {
                "professional_summary": "",
                "work_experience": [],
                "education": [],
                "continuing_studies": []
            }

print("Resume generation functions defined!")

# %%
# Cover Letter Generation Functions

def create_cover_letter_prompt(resume_data: Dict[str, Any], job_analysis: Dict[str, Any], job_title: str, company: str, job_description: str, language: str = "English") -> str:
    language_instruction = ""
    if language == "Spanish":
        language_instruction = """
        
        CRITICAL: Generate the entire cover letter in Spanish. All content, including greetings, body paragraphs, and closing must be in Spanish. Use proper Spanish business letter format and tone.
        """
    
    prompt_cover_letter = f"""
    You are an expert cover letter writer specializing in achievement-based writing and keyword optimization.{language_instruction} 

    Here is the job description:
    <job_description>
    {job_description}
    </job_description>

    Here is the candidate's resume data:
    <resume_data>
    {json.dumps(resume_data, indent=2)}
    </resume_data>

    Here is the job title:
    <job_title>
    {job_title}
    </job_title>

    Here is the company name:
    <company>
    {company}
    </company>

    Here are the relevant soft skills:
    <soft_skills>
    {json.dumps(job_analysis['soft_skills'])}
    </soft_skills>

    Here are the important keywords:
    <keywords>
    {json.dumps(job_analysis['keywords'])}
    </keywords>

    Think through your analysis internally (do not output this thinking):
    - Identify top 5 most important job requirements
    - Match soft skills to job needs and rank by importance
    - List relevant achievements from resume matching job responsibilities
    - Create skills-requirements matrix rating matches 1-5
    - Identify 3-5 current industry challenges the candidate can address
    - Review work experience timeline for accuracy
    - Find compelling career narrative or growth story
    - Brainstorm attention-grabbing opening hooks

    Now craft a cover letter following these guidelines:

    1. Structure: 3 short paragraphs (opening, body, closing) under 250 words total
    2. Tone: Casual but professional, avoid formal language or clichés
    3. Tense: Present tense (except for past experiences)
    4. Language: Avoid "expert", "skilled", "seasoned", "excited"
    5. Opening: Attention-grabbing hook
    6. Body: Highlight soft skills matching job requirements
    7. Closing: Explain how skills solve industry challenges
    8. Keywords: Naturally incorporate throughout
    9. Relevance: Only use information from resume matching job requirements
    10. Timeline: Accurately represent work experience dates

    CRITICAL: Output ONLY a valid JSON object with this exact structure. No text before or after:

    {{
        "opening_paragraph": "string (single paragraph)",
        "body_paragraph": "string (single paragraph)",
        "closing_paragraph": "string (single paragraph)"
    }}

    Do not include analysis tags, explanations, or any other text. Output pure JSON only.
    """
    return prompt_cover_letter

# Cover Letter Generation Data Validation
class CoverLetterResponse(BaseModel):
    """Model for the cover letter response structure."""
    opening_paragraph: str = "Unable to generate opening paragraph."
    body_paragraph: str = "Unable to generate body paragraph."
    closing_paragraph: str = "Unable to generate closing paragraph."
    
    @model_validator(mode='before')
    @classmethod
    def extract_from_text(cls, data):
        """Extract JSON from text if the input is not already a dict."""
        if isinstance(data, dict):
            return data
            
        if isinstance(data, str):
            # Try direct JSON parsing first
            try:
                return json.loads(data)
            except json.JSONDecodeError:
                # Extract JSON using regex if direct parsing fails
                try:
                    json_pattern = r'(\{[\s\S]*\})'
                    match = re.search(json_pattern, data)
                    if match:
                        return json.loads(match.group(1))
                except:
                    pass
                
                # If all else fails, return default content
                return {
                    "opening_paragraph": "Unable to generate opening paragraph.",
                    "body_paragraph": "Unable to generate body paragraph.",
                    "closing_paragraph": "Unable to generate closing paragraph."
                }
        
        return data

def claude_cover_letter(api_key: str, prompt_cover_letter: str) -> Dict[str, Any]:
    client = Anthropic(api_key=api_key)
    response = client.messages.create(
        model=model_name,
        max_tokens=4096,
        temperature=0.7,
        messages=[{"role": "user", "content": prompt_cover_letter}]
    )
    # Use Pydantic model to validate and extract the response
    response_text = response.content[0].text
    try:
        cover_letter_data = CoverLetterResponse.model_validate(response_text)
        return cover_letter_data.model_dump()
    except Exception as e:
        print(f"Error parsing response with Pydantic: {e}")
        # Fallback to simpler parsing if Pydantic validation fails
        try:
            return json.loads(response_text)
        except:
            print(f"JSON parsing failed. Raw response: {response_text[:150]}...")
            # Last resort: return empty data structure
            return {
                "opening_paragraph": "Unable to generate opening paragraph.",
                "body_paragraph": "Unable to generate body paragraph.",
                "closing_paragraph": "Unable to generate closing paragraph."
            }

print("Cover letter generation functions defined!")

# %%
# Document Creation Functions

def create_resume_coverletter(resume, resume_ale, cover_letter, company, country: str = "Canada", language: str = "English"):
    doc = Document()
    
    # Define section headers based on language
    if language == "Spanish":
        section_headers = {
            "experience": "EXPERIENCIA",
            "education": "EDUCACIÓN",
            "continuing_studies": "EDUCACIÓN CONTINUA"
        }
    else:
        section_headers = {
            "experience": "EXPERIENCE",
            "education": "EDUCATION",
            "continuing_studies": "CONTINUING STUDIES"
        }
    
    # Add footer with page number
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = paragraph.add_run()
    
    # Begin field char
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    
    # Field instruction text
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._element.append(instrText)
    
    # End field char
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)

    # Set footer margins
    section.footer_distance = Inches(0.3)
    
    # Set header margins
    section.header_distance = Inches(0.3)

    #----------------------------------------------#
    # GLOBAL DOCUMENT SETTINGS
    
    style = doc.styles['Normal']
    style.font.name = 'Helvetica'

    # Set margin sizes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.2)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    #----------------------------------------------#
    # RESUME

    # Name
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(resume_ale['personal_information']['full_name'])
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title.paragraph_format.space_after = Pt(1.15)
    
    # Contact info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    email_run = contact.add_run(resume_ale['personal_information']['email'])
    email_run.font.size = Pt(10)
    email_run.font.color.rgb = RGBColor(23, 54, 93)
    separator_run = contact.add_run(' | ')
    linkedin_run = contact.add_run(resume_ale['personal_information']['linkedin'])
    linkedin_run.font.size = Pt(10)
    linkedin_run.font.color.rgb = RGBColor(23, 54, 93)
    separator_run = contact.add_run(' | ')
    github_run = contact.add_run(resume_ale['personal_information']['github'])
    github_run.font.size = Pt(10)
    github_run.font.color.rgb = RGBColor(23, 54, 93)
    separator_run = contact.add_run(' | ')
    # Determine phone number based on country
    phone_number = "778.223.8536" if country == "Canada" else "33.2505.0569"
    phone_run = contact.add_run(phone_number)
    phone_run.font.size = Pt(10)
    phone_run.font.color.rgb = RGBColor(23, 54, 93)

    
    # Professional summary
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary.add_run(resume['professional_summary'])
    summary.paragraph_format.space_after = Pt(10)
    
    # Experience
    exp_header = doc.add_paragraph()
    exp_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    exp_header.add_run(section_headers["experience"]).bold = True
    exp_header.paragraph_format.space_after = Pt(0)
    
    # Set up the tab stop for right-aligned dates
    tab_stops = doc.styles['Normal'].paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.6), WD_TAB_ALIGNMENT.RIGHT)
    
    # Add work experience with right-aligned dates
    for job in resume['work_experience']:
        # Create paragraph for job header
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        
        # Add job title and company
        title_run = p.add_run(f"{job['title']} | ")
        title_run.bold = True
        company_run = p.add_run(job['company'])
        company_run.italic = True
        
        # Add right-aligned date using tab
        date_run = p.add_run('\t' + job['dates'])
        date_run.italic = True
        date_run.bold = True
        
        # Add bullet points
        for point in job['bullet_points']:
            bullet = doc.add_paragraph()
            bullet.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            bullet.paragraph_format.left_indent = Inches(0.15)
            bullet.add_run('• ' + point)
            bullet.paragraph_format.space_after = Pt(6)

    # Add Continuing Studies section
    cont_header = doc.add_paragraph()
    cont_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cont_header.add_run(section_headers["continuing_studies"]).bold = True
    cont_header.paragraph_format.space_before = Pt(10)
    cont_header.paragraph_format.space_after = Pt(0)

    # Add continuing studies entries
    for cont in resume['continuing_studies']:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        name_run = p.add_run(f"{cont['name']} | ")
        name_run.bold = True
        inst_run = p.add_run(cont['institution'])
        inst_run.italic = True
        
        # Add completion date (right-aligned)
        completion_date = p.add_run('\t' + cont['completion_date'])
        completion_date.italic = True
        completion_date.bold = True
    
    # Add Education section
    edu_header = doc.add_paragraph()
    edu_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    edu_header.add_run(section_headers["education"]).bold = True
    edu_header.paragraph_format.space_before = Pt(10)
    edu_header.paragraph_format.space_after = Pt(0)
    
    # Add education entries
    for edu in resume['education']:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        degree_run = p.add_run(f"{edu['degree']} | ")
        degree_run.bold = True
        inst_run = p.add_run(edu['institution'])
        inst_run.italic = True
        
        # Add year (right-aligned)
        completion_date = p.add_run('\t' + edu['year_completed'])
        completion_date.italic = True
        completion_date.bold = True

    # Add page break before cover letter
    if doc.paragraphs:
        last_para = doc.paragraphs[-1]
        last_para.add_run().add_break(WD_BREAK.PAGE)
    else:
        # Only if there are no paragraphs (unlikely)
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    #----------------------------------------------
    # COVER LETTER

    # Name
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(resume_ale['personal_information']['full_name'])
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title.paragraph_format.space_after = Pt(1.15)
    
    # Contact info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    email_run = contact.add_run(resume_ale['personal_information']['email'])
    email_run.font.size = Pt(10)
    email_run.font.color.rgb = RGBColor(23, 54, 93)
    separator_run = contact.add_run(' | ')
    linkedin_run = contact.add_run(resume_ale['personal_information']['linkedin'])
    linkedin_run.font.size = Pt(10)
    linkedin_run.font.color.rgb = RGBColor(23, 54, 93)
    separator_run = contact.add_run(' | ')
    github_run = contact.add_run(resume_ale['personal_information']['github'])
    github_run.font.size = Pt(10)
    github_run.font.color.rgb = RGBColor(23, 54, 93)
    separator_run = contact.add_run(' | ')
    # Use the same phone number variable as above
    phone_run = contact.add_run(phone_number)
    phone_run.font.size = Pt(10)
    phone_run.font.color.rgb = RGBColor(23, 54, 93)

    # Current Date
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date.add_run(datetime.now().strftime('%B %d, %Y'))

    # Add space after the date
    date.paragraph_format.space_after = Pt(24)
    
    # Recipient info (Company name)    
    recipient = doc.add_paragraph()
    recipient.alignment = WD_ALIGN_PARAGRAPH.LEFT
    recipient.add_run(company)
    recipient.paragraph_format.space_after = Pt(24)
    
    # Greeting
    greeting = doc.add_paragraph()
    greeting.alignment = WD_ALIGN_PARAGRAPH.LEFT
    greeting.add_run('Dear Hiring Manager:')
    greeting.paragraph_format.space_after = Pt(12)
    
    # Opening paragraph
    opening = doc.add_paragraph()
    opening.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    opening.add_run(cover_letter['opening_paragraph'])
    opening.paragraph_format.space_after = Pt(12)
    
    # Body paragraph
    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.add_run(cover_letter['body_paragraph'])
    body.paragraph_format.space_after = Pt(12)
    
    # Closing paragraph
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    closing.add_run(cover_letter['closing_paragraph'])
    closing.paragraph_format.space_after = Pt(24)
    
    # Signature
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.LEFT
    signature.add_run('Sincerely,\n\nAlejandro Leiva')
    
    return doc

def convert_word_to_pdf(file_path: str) -> str:
    pdf_path = file_path.replace('.docx', '.pdf')
    pythoncom.CoInitialize()  # Initialize COM
    convert(file_path, pdf_path)
    pythoncom.CoUninitialize()  # Uninitialize COM
    return pdf_path

print("Document creation functions defined!")

# %%
# Orchestrating Function
def create_resume(job_url: str, brightdata_api_key: str, resume_yaml_path: str = 'resume_ale.yaml') -> tuple[str, str, Dict[str, Any]]:
    """
    Orchestrate the complete resume and cover letter generation process.
    
    This function serves as the main entry point for generating tailored resumes
    and cover letters from job postings. It extracts job information, analyzes
    requirements, generates customized content using AI, and creates professional
    documents.
    
    Parameters
    ----------
    job_url : str
        URL of the job posting (LinkedIn or Indeed supported).
    brightdata_api_key : str
        BrightData API key for job extraction.
    resume_yaml_path : str, optional
        Path to the YAML file containing resume data. Default is 'resume_ale.yaml'.
    
    Returns
    -------
    tuple[str, str, Dict[str, Any]]
        A tuple containing the paths to the generated Word document and PDF, and the job dictionary.
        Format: (word_document_path, pdf_document_path, job_data_dict)
    
    Raises
    ------
    ValueError
        If required environment variables are missing or job extraction fails.
    FileNotFoundError
        If the resume YAML file is not found.
    Exception
        For any other errors during the process.
    
    Examples
    --------
    Generate documents for a LinkedIn job:
    
    >>> doc_path, pdf_path = main("https://www.linkedin.com/jobs/view/12345")
    >>> print(f"Generated: {doc_path}, {pdf_path}")
    Generated: Data_Scientist_Company_2025_09_02.docx, Data_Scientist_Company_2025_09_02.pdf
    
    Use a custom resume file:
    
    >>> doc_path, pdf_path = main("https://ca.indeed.com/viewjob?jk=67890", "my_resume.yaml")
    """
    try:
        # Step 1: Extract job information
        print("Extracting job information from URL")
        job_dict, job_df = extract_job(job_url, brightdata_api_key)
        job_title = job_dict.get('job_title', 'Unknown Title')
        company = job_dict.get('company_name', 'Unknown Company')
        job_description = job_dict.get('job_summary', '') or job_dict.get('description_text', 'No description available')
        print(f"Extracted: {job_title} at {company}")
        
        # Step 2: Load resume data
        print("Loading resume data")
        resume_data = load_resume_yaml(resume_yaml_path)
        
        # Step 3: Analyze job description
        print("Analyzing job description")
        prompt_analysis = create_analysis_prompt(job_title, company, job_description)
        if not key:
            raise ValueError("CLAUDE_API_KEY environment variable is required.")
        job_analysis = claude_analysis(key, prompt_analysis)
        
        # Step 4: Generate tailored resume
        print("Generating tailored resume")
        prompt_resume = create_resume_prompt(resume_data, job_analysis, job_title, company, job_description)
        resume = claude_resume(key, prompt_resume)
        
        # Step 5: Generate tailored cover letter
        print("Generating tailored cover letter")
        prompt_cover_letter = create_cover_letter_prompt(resume, job_analysis, job_title, company, job_description)
        cover_letter = claude_cover_letter(key, prompt_cover_letter)
        
        # Step 6: Create and save documents
        print("Creating and saving documents")
        today = datetime.today().strftime('%Y_%m_%d')
        doc = create_resume_coverletter(resume, resume_data, cover_letter, company)
        doc_path = f'Alejandro_Leiva_{job_title}_{company}_{today}.docx'
        doc.save(doc_path)
        pdf_path = convert_word_to_pdf(doc_path)
        
        return doc_path, pdf_path, job_dict
        
    except Exception as e:
        print(f"❌ Error in main process: {e}")
        raise


